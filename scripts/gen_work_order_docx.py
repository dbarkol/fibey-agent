"""
Generate a work order .docx file with the same data as work_order_fiber_splice.pdf.

Purpose: Demo contrast between CU modes:
  - None (OpenAI only)  → OpenAI rejects .docx — cannot read it
  - Basic CU            → CU extracts markdown from .docx, agent can read it
  - Classify & Analyze  → CU classifies as work_order, extracts structured fields

Expected JSON: content-understanding/demo_files/work_order_fiber_splice.json

Usage:
    uv run python scripts/gen_work_order_docx.py

Output: content-understanding/demo_files/work_order_fiber_splice.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

_REPO_ROOT = Path(__file__).parent.parent
_OUT = _REPO_ROOT / "content-understanding" / "demo_files" / "work_order_fiber_splice.docx"
_EXPECTED_JSON = _REPO_ROOT / "content-understanding" / "demo_files" / "work_order_fiber_splice.json"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def bold_run(para, text: str, size: int = 11, color: str | None = None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    return run


def normal_run(para, text: str, size: int = 11):
    run = para.add_run(text)
    run.font.size = Pt(size)
    return run


def main():
    expected = json.loads(_EXPECTED_JSON.read_text())
    doc = Document()

    # ── Title block ──────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bold_run(title_para, "FIBER SPLICE RESTORATION — SPRINGFIELD BUSINESS PARK", size=14, color="1a237e")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    normal_run(sub, f"Status: {expected['status'].upper()}  ·  Priority: {expected['priority'].upper()}  ·  Due: {expected['due_date']}", size=10)

    doc.add_paragraph()

    # ── Header metadata ─────────────────────────────────────────────────────────
    # "Field Technician" = the on-site contact the dispatched tech will meet (John Smith).
    # The LLM reads this label and returns John Smith as the assigned tech (WRONG) in Basic CU.
    # The actual dispatched technician (J. Martinez) is only in the Dispatch Log as "Route → J. Martinez".
    # The custom analyzer ignores the header and extracts from the Dispatch Log (CORRECT).
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    for i, header in enumerate(["Field Technician", "On-Site Technician", "Location", "Created"]):
        c = tbl.rows[0].cells[i]
        c.text = header
        c.paragraphs[0].runs[0].bold = True
        set_cell_bg(c, "c5cae9")
    tbl.rows[1].cells[0].text = "John Smith"
    tbl.rows[1].cells[1].text = "\u2014"           # em-dash — (none assigned)
    tbl.rows[1].cells[2].text = expected["location"]
    tbl.rows[1].cells[3].text = "2026-05-18 08:15 PDT"

    doc.add_paragraph()

    # ── Dispatch Log block (J. Martinez appears as routing reference, NOT labeled "technician") ──
    # Demo intent: LLM reading flat markdown sees "Field Technical Contact: Marcus Tran"
    # as the prominent labeled name and returns that (WRONG). The custom analyzer description
    # says to look for the name after "Route →" in the Dispatch Log — so it returns J. Martinez (correct).
    disp = doc.add_table(rows=1, cols=2)
    disp.style = "Table Grid"
    label_cell = disp.rows[0].cells[0]
    label_cell.text = "Dispatch Log"
    label_cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(label_cell, "dce3f0")
    disp.rows[0].cells[1].text = (
        f"2026-05-18 08:15 PDT  |  NOC Ref: WO-DISP-0518  |  "
        f"Dispatcher: R. Singh  |  Route \u2192 {expected['assigned_technician']}  |  Status: Pending Accept"
    )
    set_cell_bg(disp.rows[0].cells[1], "dce3f0")

    doc.add_paragraph()

    # ── Job description ───────────────────────────────────────────────────────
    add_heading(doc, "Job Description", level=2)
    doc.add_paragraph(expected["description"])

    doc.add_paragraph()

    # ── Site access & contact ─────────────────────────────────────────────────
    add_heading(doc, "Site Access & Contact", level=2)
    p = doc.add_paragraph()
    bold_run(p, "Contact: ", size=11)
    normal_run(p, "John Smith — Network Operations Supervisor  |  (425) 555-0183", size=11)
    p2 = doc.add_paragraph()
    bold_run(p2, "Access Notes: ", size=11)
    normal_run(p2, "Escort required. Check in at security desk (lobby entrance). Hard hat + safety vest mandatory in parking structure.", size=11)

    doc.add_paragraph()

    # ── Safety protocols ─────────────────────────────────────────────────────
    add_heading(doc, "Safety Protocols", level=2)
    for item in [
        "Wear PPE: safety glasses, cut-resistant gloves, and high-vis vest",
        "Lock out / tag out power to cabinet before opening enclosure",
        "Confirm no live laser sources before handling bare fiber",
        "Dispose of fiber cleave scraps in designated sharps container",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_paragraph()

    # ── Parts & materials ─────────────────────────────────────────────────────
    add_heading(doc, "Parts & Materials Required", level=2)
    parts_tbl = doc.add_table(rows=1, cols=3)
    parts_tbl.style = "Table Grid"
    for i, label in enumerate(["Part ID", "Description", "Qty"]):
        c = parts_tbl.rows[0].cells[i]
        c.text = label
        c.paragraphs[0].runs[0].bold = True
        set_cell_bg(c, "c5cae9")

    part_descriptions = {
        "FIB-003": "Single-Mode Splice Tray (12-fiber)",
        "FIB-012": "Fiber Splice Enclosure (24-count)",
    }
    for part in expected["parts_needed"]:
        row = parts_tbl.add_row()
        row.cells[0].text = part["part_id"]
        row.cells[1].text = part_descriptions.get(part["part_id"], "")
        row.cells[2].text = str(part["quantity"])

    doc.add_paragraph()

    # ── Completion checklist ──────────────────────────────────────────────────
    add_heading(doc, "Field Completion Checklist", level=2)
    for item in [
        "Review as-built drawings for cable route",
        "Confirm part availability",
        "Notify NOC of maintenance window",
        "Perform fusion splices (target IL ≤ 0.10 dB each)",
        "Install new enclosure and reseal per manufacturer spec",
        "Run end-to-end OTDR trace on all 24 fibers",
        "Confirm signal restore for all 14 affected tenant circuits",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_paragraph()

    # ── Sign-Off (Signature blank — OPEN; Print Name pre-filled with assigned tech) ─────────────
    add_heading(doc, "Sign-Off & Completion", level=2)
    sign_tbl = doc.add_table(rows=2, cols=4)
    sign_tbl.style = "Table Grid"
    for i, label in enumerate(["Role", "Signature", "Print Name", "Date"]):
        c = sign_tbl.rows[0].cells[i]
        c.text = label
        c.paragraphs[0].runs[0].bold = True
        set_cell_bg(c, "c5cae9")
    sign_tbl.rows[1].cells[0].text = "Technician Signature"
    sign_tbl.rows[1].cells[1].text = ""   # Signature blank
    sign_tbl.rows[1].cells[2].text = ""   # Print Name blank — J. Martinez only in Dispatch Log
    sign_tbl.rows[1].cells[3].text = ""

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    normal_run(footer, "Fibey Field Ops  ·  Fiber Splice Restoration  ·  CONFIDENTIAL", size=9)

    _OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(_OUT)
    print(f"Saved: {_OUT}")


if __name__ == "__main__":
    main()
